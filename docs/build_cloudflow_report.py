from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "Отчет_CloudFlow_Marketplace_исправленный.docx"
TMP = BASE / "_cloudflow_report_assets"
TMP.mkdir(parents=True, exist_ok=True)


def font(size=28, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size=size)
    return ImageFont.load_default()


FONT = font(28)
BOLD = font(30, True)
SMALL = font(22)
SMALL_BOLD = font(23, True)


def wrap_text(text, draw, max_width, fnt):
    lines = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        line = ""
        for word in words:
            candidate = f"{line} {word}".strip()
            width = draw.textbbox((0, 0), candidate, font=fnt)[2]
            if width <= max_width or not line:
                line = candidate
            else:
                lines.append(line)
                line = word
        if line:
            lines.append(line)
    return lines


def centered_box(draw, box, title, subtitle="", fill="#F8FAFC", outline="#334155"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    max_width = x2 - x1 - 32
    title_lines = wrap_text(title, draw, max_width, BOLD)
    sub_lines = wrap_text(subtitle, draw, max_width, SMALL) if subtitle else []
    total = len(title_lines) * 36 + len(sub_lines) * 28 + (10 if sub_lines else 0)
    y = y1 + ((y2 - y1) - total) / 2
    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=BOLD)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=BOLD, fill="#0F172A")
        y += 36
    if sub_lines:
        y += 6
    for line in sub_lines:
        bbox = draw.textbbox((0, 0), line, font=SMALL)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=SMALL, fill="#334155")
        y += 28


def arrow(draw, start, end, color="#475569", width=5):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    import math

    angle = math.atan2(ey - sy, ex - sx)
    size = 16
    p1 = (ex - size * math.cos(angle - math.pi / 7), ey - size * math.sin(angle - math.pi / 7))
    p2 = (ex - size * math.cos(angle + math.pi / 7), ey - size * math.sin(angle + math.pi / 7))
    draw.polygon([end, p1, p2], fill=color)


def make_architecture_diagram(path):
    img = Image.new("RGB", (1800, 1120), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((60, 36), "Архитектура CloudFlow Marketplace", font=font(42, True), fill="#0F172A")
    boxes = {
        "user": (70, 170, 390, 310),
        "next": (520, 150, 900, 330),
        "routes": (1040, 150, 1420, 330),
        "go": (520, 510, 900, 720),
        "db": (1080, 500, 1450, 700),
        "files": (1080, 760, 1450, 910),
        "smtp": (160, 530, 390, 690),
    }
    centered_box(d, boxes["user"], "Пользователь", "Каталог, профиль,\nпокупка, публикация")
    centered_box(d, boxes["next"], "Next.js 15 + React 19", "App Router, страницы,\nкомпоненты UI")
    centered_box(d, boxes["routes"], "API-маршруты Next.js", "Прокси к GraphQL,\nсессия, загрузка файлов")
    centered_box(d, boxes["go"], "Go backend", "Gin, gqlgen GraphQL,\nJWT, бизнес-логика")
    centered_box(d, boxes["db"], "PostgreSQL", "users, products,\norders, server_access")
    centered_box(d, boxes["files"], "Файловое хранилище", "Изображения товаров\nи аватары")
    centered_box(d, boxes["smtp"], "SMTP-сервис", "Коды email 2FA")
    arrow(d, (390, 240), (520, 240))
    arrow(d, (900, 240), (1040, 240))
    arrow(d, (1230, 330), (810, 510))
    arrow(d, (900, 610), (1080, 600))
    arrow(d, (1230, 760), (1230, 700))
    arrow(d, (520, 610), (390, 610))
    d.text((420, 205), "HTTP", font=SMALL_BOLD, fill="#334155")
    d.text((930, 205), "JSON", font=SMALL_BOLD, fill="#334155")
    d.text((934, 390), "GraphQL /query", font=SMALL_BOLD, fill="#334155")
    d.text((925, 570), "SQL через pgx", font=SMALL_BOLD, fill="#334155")
    d.text((402, 575), "email code", font=SMALL_BOLD, fill="#334155")
    d.text((80, 1010), "Роли: User, Seller, Moderator, Creator. Закрытые операции выполняются только с JWT Bearer token.", font=SMALL, fill="#475569")
    img.save(path, quality=95)


def make_er_diagram(path):
    img = Image.new("RGB", (1800, 1180), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((60, 36), "Логическая модель данных", font=font(42, True), fill="#0F172A")
    boxes = {
        "users": (90, 170, 430, 330),
        "products": (720, 160, 1100, 350),
        "categories": (1320, 145, 1640, 285),
        "images": (1320, 330, 1640, 470),
        "orders": (90, 610, 430, 770),
        "order_items": (720, 600, 1100, 790),
        "server_access": (1320, 565, 1640, 725),
        "tags": (1320, 770, 1640, 930),
        "reviews": (720, 880, 1100, 1040),
    }
    fills = {
        "users": "#EFF6FF",
        "products": "#F0FDFA",
        "categories": "#F8FAFC",
        "images": "#F8FAFC",
        "orders": "#FFF7ED",
        "order_items": "#FEFCE8",
        "server_access": "#FDF2F8",
        "tags": "#F5F3FF",
        "reviews": "#ECFDF5",
    }
    labels = {
        "users": "users\nid, email, role,\nbalance, rating",
        "products": "products\nseller_id, category_id,\nprice, status",
        "categories": "categories\nparent_id",
        "images": "images\ntarget_id, file_name,\nis_preview",
        "orders": "orders\nbuyer_id, status,\ntotal_amount",
        "order_items": "order_items\norder_id, product_id,\nseller_id, price",
        "server_access": "server_access\nproduct_id, ip,\nssh login/key",
        "tags": "tags + product_tag\nхарактеристики сервера",
        "reviews": "seller_reviews\nseller_id, buyer_id,\nproduct_id, rating",
    }
    for key, box in boxes.items():
        centered_box(d, box, labels[key], "", fill=fills[key], outline="#475569")

    def edge(key, side):
        x1, y1, x2, y2 = boxes[key]
        return {
            "left": (x1, (y1 + y2) // 2),
            "right": (x2, (y1 + y2) // 2),
            "top": ((x1 + x2) // 2, y1),
            "bottom": ((x1 + x2) // 2, y2),
        }[side]

    def relation(start, end, label, offset=(0, 0)):
        arrow(d, start, end, color="#64748B", width=4)
        mx, my = (start[0] + end[0]) // 2 + offset[0], (start[1] + end[1]) // 2 + offset[1]
        d.rounded_rectangle((mx - 95, my - 17, mx + 95, my + 17), radius=8, fill="#FFFFFF", outline="#CBD5E1")
        tb = d.textbbox((0, 0), label, font=font(17, True))
        d.text((mx - (tb[2] - tb[0]) / 2, my - 11), label, font=font(17, True), fill="#334155")

    relation(edge("users", "right"), edge("products", "left"), "seller_id")
    relation(edge("categories", "left"), edge("products", "right"), "category_id", (0, -20))
    relation(edge("products", "right"), edge("images", "left"), "target_id")
    relation(edge("users", "bottom"), edge("orders", "top"), "buyer_id")
    relation(edge("orders", "right"), edge("order_items", "left"), "order_id")
    relation(edge("products", "bottom"), edge("order_items", "top"), "product_id", (110, 0))
    relation(edge("products", "right"), edge("server_access", "left"), "product_id", (0, 18))
    relation(edge("products", "right"), edge("tags", "left"), "product_tag", (0, 8))
    relation(edge("order_items", "bottom"), edge("reviews", "top"), "после покупки")

    d.text((80, 1100), "server_access отделяет публичную карточку товара от конфиденциальных параметров подключения.", font=SMALL, fill="#475569")
    img.save(path, quality=95)


def make_purchase_diagram(path):
    img = Image.new("RGB", (1800, 1050), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((60, 36), "Сценарий покупки и выдачи доступа", font=font(42, True), fill="#0F172A")
    steps = [
        ("1", "Вход и 2FA", "Логин, пароль,\nemail-код и JWT"),
        ("2", "Выбор сервера", "Активные публичные\nкарточки каталога"),
        ("3", "Покупка", "Проверка баланса,\nстатуса и прав"),
        ("4", "Транзакция", "Заказ, позиция,\nсписание баланса"),
        ("5", "Доступ", "server_access виден\nпокупателю"),
    ]
    x = 90
    y = 235
    w = 285
    h = 290
    for idx, (num, title, text) in enumerate(steps):
        box = (x + idx * 335, y, x + idx * 335 + w, y + h)
        d.rounded_rectangle(box, radius=22, fill="#F8FAFC", outline="#334155", width=3)
        d.ellipse((box[0] + 22, box[1] + 22, box[0] + 78, box[1] + 78), fill="#2563EB")
        tb = d.textbbox((0, 0), num, font=BOLD)
        d.text((box[0] + 50 - (tb[2] - tb[0]) / 2, box[1] + 33), num, font=BOLD, fill="#FFFFFF")
        d.text((box[0] + 26, box[1] + 100), title, font=BOLD, fill="#0F172A")
        ty = box[1] + 152
        for line in wrap_text(text, d, w - 48, SMALL):
            d.text((box[0] + 26, ty), line, font=SMALL, fill="#334155")
            ty += 30
        if idx < len(steps) - 1:
            arrow(d, (box[2] + 16, y + h // 2), (box[2] + 48, y + h // 2), color="#475569", width=5)
    centered_box(d, (360, 690, 780, 865), "Контроль доступа", "JWT + роль + принадлежность\nтовара покупателю/продавцу", fill="#ECFDF5")
    centered_box(d, (1020, 690, 1440, 865), "Асинхронность", "HTTP-запросы обрабатываются\nпараллельно; email-коды имеют TTL", fill="#FFF7ED")
    arrow(d, (780, 775), (1020, 775), color="#64748B", width=4)
    d.text((80, 965), "Закрытые IP, логин, пароль, ключ и примечания не попадают в публичный каталог.", font=SMALL, fill="#475569")
    img.save(path, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_run_font(run, size=14, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, 14)
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_heading(doc, number, title, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    text = f"{number} {title}" if number else title
    run = p.add_run(text)
    set_run_font(run, 14 if level > 1 else 16, True)
    return p


def add_caption(doc, text, figure=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def add_table(doc, caption, headers, rows, widths):
    add_caption(doc, caption, figure=False)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        set_cell_text(hdr[i], head, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
        hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].width = widths[i]
    doc.add_paragraph()
    return table


make_architecture_diagram(TMP / "architecture.png")
make_er_diagram(TMP / "er.png")
make_purchase_diagram(TMP / "purchase.png")

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
styles["Normal"].font.size = Pt(14)

for style_name in ["Heading 1", "Heading 2", "Title"]:
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

footer = section.footer.paragraphs[0]
add_page_number(footer)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(18)
run = title.add_run("Отчет по лабораторным работам\nпо курсу «Расширенные возможности языков высокого уровня»")
set_run_font(run, 16, True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
run = subtitle.add_run("Проект: CloudFlow Marketplace для публикации и покупки серверной инфраструктуры")
set_run_font(run, 15, True)

meta = [
    ("Дисциплина", "Расширенные возможности языков высокого уровня"),
    ("Учебный год / семестр", "2025/2026, весна"),
    ("Преподаватель", "Борисов Артем Ильич, ассистент кафедры 304"),
    ("Формат отчета", "Итоговый отчет без вставки исходного кода"),
]
add_table(doc, "Таблица 1 - Сведения об отчете", ["Параметр", "Значение"], meta, [Cm(5), Cm(11)])

add_heading(doc, "1", "Описание темы проекта")
for text in [
    "Темой проекта является разработка торговой площадки для продажи и покупки серверных ресурсов. В качестве товара выступают реальные серверные предложения: выделенные серверы, VPS и хостинг-узлы с понятными характеристиками, такими как количество ядер процессора, объем оперативной памяти, дисковое пространство, регион размещения, пропускная способность и операционная система.",
    "Актуальность проекта связана с тем, что типовые доски объявлений не дают удобного механизма для безопасной публикации технических параметров, приема оплаты, хранения приватных данных доступа и разграничения прав между продавцом, покупателем и администратором. В CloudFlow Marketplace эта проблема решается за счет разделения публичной карточки сервера и защищенного блока с конфиденциальной информацией, который становится доступен только после подтвержденной покупки.",
    "На этапе планирования предполагалось реализовать регистрацию и вход пользователей, публикацию карточек серверов, просмотр каталога с категориями, покупку серверов за внутренний баланс, отображение купленных серверов в профиле пользователя, отправку дополнительного проверочного кода при входе, хранение изображений и аватаров, а также административное управление разделами маркетплейса.",
]:
    add_paragraph(doc, text)

add_table(
    doc,
    "Таблица 2 - Соответствие проекта требованиям лабораторного задания",
    ["Требование", "Реализация в CloudFlow Marketplace"],
    [
        ("Пользовательский интерфейс", "Next.js 15, React 19, TypeScript, Tailwind CSS, shadcn/ui-компоненты."),
        ("Сетевое взаимодействие", "HTTP API-маршруты Next.js проксируют запросы к Go backend; основной контракт данных описан GraphQL."),
        ("База данных", "PostgreSQL хранит пользователей, товары, категории, заказы, отзывы, изображения, теги и закрытые параметры доступа к серверам."),
        ("Авторизация, аутентификация, 2FA", "JWT используется для защищенных операций; email 2FA реализована через одноразовый код и SMTP."),
        ("Асинхронность и параллельная работа", "Запросы интерфейса выполняются асинхронно; Go/Gin обрабатывает HTTP-запросы параллельно, а pgx использует подключение к PostgreSQL."),
        ("Функциональные модули", "Публикация серверов, каталог, покупка, профиль, администрирование, отзывы продавцам и выдача закрытых данных после покупки."),
    ],
    [Cm(5.2), Cm(10.8)],
)

add_heading(doc, "2", "Описание проектных и архитектурных решений")
for text in [
    "При разработке проекта была выбрана многослойная архитектура с явным разделением клиентской части, серверного API и слоя хранения данных. Пользовательский интерфейс реализован на Next.js 15 с использованием TypeScript, App Router и серверных маршрутов. Такой подход позволяет объединить современный интерактивный интерфейс с контролируемой работой на стороне сервера, а также скрыть детали взаимодействия с основным backend от браузера.",
    "Для построения интерфейса использованы Tailwind CSS и набор компонентов shadcn/ui. Это позволило создать единый визуальный стиль, сократить объем повторяющегося кода и упростить поддержку форм, диалоговых окон, выпадающих меню и навигации. Отдельное внимание было уделено модальному окну авторизации, страницам маркетплейса, публикации серверов и пользовательскому профилю.",
    "Серверная часть построена на языке Go. В качестве HTTP-уровня используется Gin, а основным контрактом обмена данными служит GraphQL, реализованный через gqlgen. Такой выбор удобен тем, что frontend получает строго описанные структуры, а backend централизованно описывает запросы, мутации и правила доступа. Для хранения данных используется PostgreSQL.",
    "Для аутентификации используется JWT. После регистрации и входа пользователь получает токен, который применяется при дальнейших обращениях к защищенным операциям. Дополнительно реализован сценарий двухфакторной проверки по электронной почте: после ввода логина и пароля backend генерирует одноразовый код и отправляет его через SMTP.",
    "Отдельным проектным решением стало разделение открытых и закрытых данных товара. В публичной карточке сервера хранятся название, описание, цена, категория, характеристики и изображение. Конфиденциальные сведения, такие как IP-адрес, SSH-логин, пароль, приватный ключ и служебные примечания, вынесены в отдельную таблицу server_access.",
]:
    add_paragraph(doc, text)

add_table(
    doc,
    "Таблица 3 - Основные технологии и причины выбора",
    ["Компонент", "Технологии", "Обоснование"],
    [
        ("Frontend", "Next.js 15, React 19, TypeScript", "Поддержка современных маршрутов, компонентного UI и типизации клиентской части."),
        ("UI-слой", "Tailwind CSS, shadcn/ui, lucide-react", "Единый стиль, готовые элементы форм и навигации, быстрая доработка интерфейса."),
        ("Backend", "Go, Gin, gqlgen", "Производительная обработка HTTP-запросов, строгий GraphQL-контракт и удобная генерация серверного слоя."),
        ("Данные", "PostgreSQL, pgx", "Реляционная модель подходит для пользователей, заказов, товаров и связей между ними."),
        ("Безопасность", "JWT, email OTP, SMTP", "Разграничение защищенных операций и дополнительная проверка личности при входе."),
    ],
    [Cm(3.2), Cm(5.1), Cm(7.5)],
)

add_heading(doc, "3", "Описание конечной архитектуры системы")
add_paragraph(
    doc,
    "Итоговая архитектура проекта состоит из нескольких взаимодействующих уровней. На клиентском уровне находится пользовательский интерфейс маркетплейса. Он отвечает за отображение карточек серверов, формы публикации, профиль пользователя, экран покупки, административные разделы и обработку пользовательских действий. На промежуточном уровне расположены серверные маршруты Next.js, которые принимают запросы от интерфейса, проверяют пользовательскую сессию и проксируют обращения к backend. На основном серверном уровне работает GraphQL API на Go, содержащее бизнес-логику, авторизацию и взаимодействие с базой данных."
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run()
run.add_picture(str(TMP / "architecture.png"), width=Inches(6.7))
add_caption(doc, "Рисунок 1 - Схема модулей и взаимодействия компонентов CloudFlow Marketplace")

add_paragraph(
    doc,
    "В базе данных система опирается на несколько ключевых сущностей. Пользователи содержат учетные данные, баланс, рейтинг продавца и ссылку на аватар. Категории определяют структуру разделов маркетплейса. Продукты описывают опубликованные серверы. Теги и промежуточная таблица связей хранят технические характеристики сервера. Заказы и позиции заказов фиксируют факт покупки. Отдельная таблица с приватными параметрами доступа хранит конфиденциальные сведения о сервере. Также предусмотрены отзывы продавцам, на основании которых пересчитывается их рейтинг."
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run()
run.add_picture(str(TMP / "er.png"), width=Inches(6.7))
add_caption(doc, "Рисунок 2 - Логическая схема основных сущностей базы данных")

add_paragraph(
    doc,
    "После покупки товара система изменяет его статус, чтобы предложение исчезало из общего списка доступных к покупке серверов. Одновременно заказ связывается с покупателем, и этот пользователь начинает видеть сервер в своем профиле. Именно там он получает доступ к приватным данным подключения. Таким образом обеспечивается логика «один товар - один покупатель» для серверов, которые продаются как уникальные позиции."
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run()
run.add_picture(str(TMP / "purchase.png"), width=Inches(6.7))
add_caption(doc, "Рисунок 3 - Сценарий покупки сервера и выдачи закрытых данных")

add_heading(doc, "4", "Описание основных модулей и функций")
modules = [
    ("Модуль аутентификации и регистрации", "Включает формы входа и создания аккаунта, проверку учетных данных, получение JWT, выход из системы и двухшаговую почтовую верификацию. От этого модуля зависит доступ ко всем защищенным операциям маркетплейса."),
    ("Модуль маркетплейса", "Отображает список опубликованных серверов, категории, карточки товаров и детальную страницу каждого предложения. Пользователь анализирует характеристики, стоимость, рейтинг продавца и принимает решение о покупке."),
    ("Центр публикации серверов", "Позволяет продавцу создать карточку товара, указать описание сервера, стоимость, категорию, характеристики оборудования и приватные данные подключения. Также поддерживается редактирование опубликованных позиций и замена изображения."),
    ("Профиль пользователя", "Содержит редактирование учетных данных, загрузку аватара, пополнение тестового баланса, просмотр купленных серверов и доступ к закрытым параметрам подключения. В профиле также отображаются отзывы и рейтинг продавца."),
    ("Административная панель", "Используется для управления пользователями, категориями и карточками маркетплейса. Наличие отдельного административного интерфейса показывает, что система проектируется как расширяемое приложение с отдельным уровнем управления."),
    ("Серверный модуль бизнес-логики", "Реализует GraphQL-запросы и мутации, работу с базой данных, обработку покупки, начисление средств продавцу, списание баланса покупателя, сохранение отзывов и выдачу доступа к приватным данным после покупки."),
]
for title_text, body in modules:
    add_heading(doc, None, title_text, level=2)
    add_paragraph(doc, body)

add_heading(doc, "5", "Заключение")
add_paragraph(
    doc,
    "В результате выполнения лабораторных работ был разработан учебный, но практически ориентированный прототип серверного маркетплейса CloudFlow. Проект демонстрирует применение современных средств frontend- и backend-разработки, работу с типизированным API, организацию безопасной аутентификации, хранение приватных данных, построение ролевой модели и реализацию бизнес-процесса покупки серверной инфраструктуры. Структура решения позволяет расширять проект в дальнейшем, например, за счет реальной платежной интеграции, поддержки нескольких экземпляров товара, расширенной аналитики и более глубокого администрирования."
)

doc.save(OUT)
print(OUT)
