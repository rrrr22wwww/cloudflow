from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from build_cloudflow_report import (
    TMP,
    make_architecture_diagram,
    make_er_diagram,
    make_purchase_diagram,
)


BASE = Path(__file__).resolve().parent
OUT = BASE / "Отчет_CloudFlow_Marketplace_по_оформлению_методички.docx"


def set_run_font(run, size=14, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def setup_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.different_first_page_header_footer = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, 12)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_center(doc, text="", size=14, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if text:
        run = p.add_run(text)
        set_run_font(run, size, bold)
    return p


def add_body(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        set_run_font(run, 14)
    return p


def add_heading(doc, text, numbered=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text.upper() if not numbered else text)
    set_run_font(run, 14, True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 14, True)
    return p


def add_caption(doc, text, table=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if table else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(1.5 if table else 0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, 12, bold)
    for margin in ("top", "bottom", "left", "right"):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def add_table(doc, caption, headers, rows, widths):
    add_caption(doc, caption, table=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_text(cell, header, True, WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = widths[i]
            set_cell_text(cells[i], value)
    return table


def add_toc_row(doc, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    dots = "." * max(6, 72 - len(title))
    run = p.add_run(f"{title} {dots} {page}")
    set_run_font(run, 14)


def build():
    TMP.mkdir(exist_ok=True)
    make_architecture_diagram(TMP / "architecture.png")
    make_er_diagram(TMP / "er.png")
    make_purchase_diagram(TMP / "purchase.png")

    doc = Document()
    setup_section(doc.sections[0])
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(14)

    add_center(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", 12, False, after=6)
    add_center(doc, "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ\nУЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", 12, False, after=6)
    add_center(doc, "«МОСКОВСКИЙ АВИАЦИОННЫЙ ИНСТИТУТ\n(национальный исследовательский университет)»", 12, False, after=28)
    add_center(doc, "Кафедра 304", 12, False, after=90)
    add_center(doc, "ОТЧЕТ", 16, True, after=4)
    add_center(doc, "по лабораторным работам", 14, True, after=6)
    add_center(doc, "по курсу «Расширенные возможности языков высокого уровня»", 14, False, after=20)
    add_center(doc, "На тему: CloudFlow Marketplace для публикации и покупки серверной инфраструктуры", 14, False, after=55)
    add_body(doc, "Отчет выполнил ________________________________ (________________)")
    add_body(doc, "Отчет принял  ________________________________ (________________)")
    add_center(doc, "Москва, 2026", 14, True, before=115)

    page_break(doc)
    add_heading(doc, "СОДЕРЖАНИЕ", numbered=False)
    toc = [
        ("1 Описание темы проекта", 3),
        ("2 Описание проектных и архитектурных решений", 5),
        ("3 Описание конечной архитектуры системы", 7),
        ("4 Описание основных модулей и функций", 10),
        ("5 Заключение", 12),
    ]
    for title, num in toc:
        add_toc_row(doc, title, num)

    page_break(doc)
    add_heading(doc, "1 ОПИСАНИЕ ТЕМЫ ПРОЕКТА")
    for text in [
        "Темой проекта является разработка торговой площадки для продажи и покупки серверных ресурсов. В качестве товара выступают реальные серверные предложения: выделенные серверы, VPS и хостинг-узлы с понятными характеристиками, такими как количество ядер процессора, объем оперативной памяти, дисковое пространство, регион размещения, пропускная способность и операционная система.",
        "Актуальность проекта связана с тем, что типовые доски объявлений не дают удобного механизма для безопасной публикации технических параметров, приема оплаты, хранения приватных данных доступа и разграничения прав между продавцом, покупателем и администратором. В CloudFlow Marketplace эта проблема решается за счет разделения публичной карточки сервера и защищенного блока с конфиденциальной информацией, который становится доступен только после подтвержденной покупки.",
        "На этапе планирования предполагалось реализовать регистрацию и вход пользователей, публикацию карточек серверов, просмотр каталога с категориями, покупку серверов за внутренний баланс, отображение купленных серверов в профиле пользователя, отправку дополнительного проверочного кода при входе, хранение изображений и аватаров, а также административное управление разделами маркетплейса.",
    ]:
        add_body(doc, text)
    page_break(doc)
    add_table(
        doc,
        "Таблица 1.1 - Соответствие проекта требованиям лабораторного задания",
        ["Требование", "Реализация в CloudFlow Marketplace"],
        [
            ("Пользовательский интерфейс", "Next.js 15, React 19, TypeScript, Tailwind CSS, shadcn/ui-компоненты."),
            ("Сетевое взаимодействие", "HTTP API-маршруты Next.js проксируют запросы к Go backend; основной контракт данных описан GraphQL."),
            ("База данных", "PostgreSQL хранит пользователей, товары, категории, заказы, отзывы, изображения, теги и закрытые параметры доступа к серверам."),
            ("Авторизация, аутентификация, 2FA", "JWT используется для защищенных операций; email 2FA реализована через одноразовый код и SMTP."),
            ("Асинхронность и параллельная работа", "Запросы интерфейса выполняются асинхронно; Go/Gin обрабатывает HTTP-запросы параллельно, а pgx использует подключение к PostgreSQL."),
            ("Функциональные модули", "Публикация серверов, каталог, покупка, профиль, администрирование, отзывы продавцам и выдача закрытых данных после покупки."),
        ],
        [Cm(5.1), Cm(10.9)],
    )

    page_break(doc)
    add_heading(doc, "2 ОПИСАНИЕ ПРОЕКТНЫХ И АРХИТЕКТУРНЫХ РЕШЕНИЙ")
    for text in [
        "При разработке проекта была выбрана многослойная архитектура с явным разделением клиентской части, серверного API и слоя хранения данных. Пользовательский интерфейс реализован на Next.js 15 с использованием TypeScript, App Router и серверных маршрутов. Такой подход позволяет объединить современный интерактивный интерфейс с контролируемой работой на стороне сервера, а также скрыть детали взаимодействия с основным backend от браузера.",
        "Для построения интерфейса использованы Tailwind CSS и набор компонентов shadcn/ui. Это позволило создать единый визуальный стиль, сократить объем повторяющегося кода и упростить поддержку форм, диалоговых окон, выпадающих меню и навигации. Отдельное внимание было уделено модальному окну авторизации, страницам маркетплейса, публикации серверов и пользовательскому профилю.",
        "Серверная часть построена на языке Go. В качестве HTTP-уровня используется Gin, а основным контрактом обмена данными служит GraphQL, реализованный через gqlgen. Такой выбор удобен тем, что frontend получает строго описанные структуры, а backend централизованно описывает запросы, мутации и правила доступа. Для хранения данных используется PostgreSQL.",
        "Для аутентификации используется JWT. После регистрации и входа пользователь получает токен, который применяется при дальнейших обращениях к защищенным операциям. Дополнительно реализован сценарий двухфакторной проверки по электронной почте: после ввода логина и пароля backend генерирует одноразовый код и отправляет его через SMTP.",
        "Отдельным проектным решением стало разделение открытых и закрытых данных товара. В публичной карточке сервера хранятся название, описание, цена, категория, характеристики и изображение. Конфиденциальные сведения, такие как IP-адрес, SSH-логин, пароль, приватный ключ и служебные примечания, вынесены в отдельную таблицу server_access.",
    ]:
        add_body(doc, text)
    page_break(doc)
    add_table(
        doc,
        "Таблица 2.1 - Основные технологии и причины выбора",
        ["Компонент", "Технологии", "Обоснование"],
        [
            ("Frontend", "Next.js 15, React 19, TypeScript", "Поддержка современных маршрутов, компонентного UI и типизации клиентской части."),
            ("UI-слой", "Tailwind CSS, shadcn/ui, lucide-react", "Единый стиль, готовые элементы форм и навигации, быстрая доработка интерфейса."),
            ("Backend", "Go, Gin, gqlgen", "Производительная обработка HTTP-запросов, строгий GraphQL-контракт и удобная генерация серверного слоя."),
            ("Данные", "PostgreSQL, pgx", "Реляционная модель подходит для пользователей, заказов, товаров и связей между ними."),
            ("Безопасность", "JWT, email OTP, SMTP", "Разграничение защищенных операций и дополнительная проверка личности при входе."),
        ],
        [Cm(3.3), Cm(5.0), Cm(7.7)],
    )

    page_break(doc)
    add_heading(doc, "3 ОПИСАНИЕ КОНЕЧНОЙ АРХИТЕКТУРЫ СИСТЕМЫ")
    add_body(doc, "Итоговая архитектура проекта состоит из нескольких взаимодействующих уровней. На клиентском уровне находится пользовательский интерфейс маркетплейса. Он отвечает за отображение карточек серверов, формы публикации, профиль пользователя, экран покупки, административные разделы и обработку пользовательских действий. На промежуточном уровне расположены серверные маршруты Next.js, которые принимают запросы от интерфейса, проверяют пользовательскую сессию и проксируют обращения к backend. На основном серверном уровне работает GraphQL API на Go, содержащее бизнес-логику, авторизацию и взаимодействие с базой данных.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(TMP / "architecture.png"), width=Inches(6.7))
    add_caption(doc, "Рисунок 3.1 - Схема модулей и взаимодействия компонентов CloudFlow Marketplace")
    add_body(doc, "В базе данных система опирается на несколько ключевых сущностей. Пользователи содержат учетные данные, баланс, рейтинг продавца и ссылку на аватар. Категории определяют структуру разделов маркетплейса. Продукты описывают опубликованные серверы. Теги и промежуточная таблица связей хранят технические характеристики сервера. Заказы и позиции заказов фиксируют факт покупки. Отдельная таблица с приватными параметрами доступа хранит конфиденциальные сведения о сервере. Также предусмотрены отзывы продавцам, на основании которых пересчитывается их рейтинг.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(TMP / "er.png"), width=Inches(6.7))
    add_caption(doc, "Рисунок 3.2 - Логическая схема основных сущностей базы данных")
    add_body(doc, "После покупки товара система изменяет его статус, чтобы предложение исчезало из общего списка доступных к покупке серверов. Одновременно заказ связывается с покупателем, и этот пользователь начинает видеть сервер в своем профиле. Именно там он получает доступ к приватным данным подключения. Таким образом обеспечивается логика «один товар - один покупатель» для серверов, которые продаются как уникальные позиции.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(TMP / "purchase.png"), width=Inches(6.7))
    add_caption(doc, "Рисунок 3.3 - Сценарий покупки сервера и выдачи закрытых данных")

    page_break(doc)
    add_heading(doc, "4 ОПИСАНИЕ ОСНОВНЫХ МОДУЛЕЙ И ФУНКЦИЙ")
    modules = [
        ("4.1 Модуль аутентификации и регистрации", "Включает формы входа и создания аккаунта, проверку учетных данных, получение JWT, выход из системы и двухшаговую почтовую верификацию. От этого модуля зависит доступ ко всем защищенным операциям маркетплейса."),
        ("4.2 Модуль маркетплейса", "Отображает список опубликованных серверов, категории, карточки товаров и детальную страницу каждого предложения. Пользователь анализирует характеристики, стоимость, рейтинг продавца и принимает решение о покупке."),
        ("4.3 Центр публикации серверов", "Позволяет продавцу создать карточку товара, указать описание сервера, стоимость, категорию, характеристики оборудования и приватные данные подключения. Также поддерживается редактирование опубликованных позиций и замена изображения."),
        ("4.4 Профиль пользователя", "Содержит редактирование учетных данных, загрузку аватара, пополнение тестового баланса, просмотр купленных серверов и доступ к закрытым параметрам подключения. В профиле также отображаются отзывы и рейтинг продавца."),
        ("4.5 Административная панель", "Используется для управления пользователями, категориями и карточками маркетплейса. Наличие отдельного административного интерфейса показывает, что система проектируется как расширяемое приложение с отдельным уровнем управления."),
        ("4.6 Серверный модуль бизнес-логики", "Реализует GraphQL-запросы и мутации, работу с базой данных, обработку покупки, начисление средств продавцу, списание баланса покупателя, сохранение отзывов и выдачу доступа к приватным данным после покупки."),
    ]
    for title, body in modules:
        add_subheading(doc, title)
        add_body(doc, body)

    page_break(doc)
    add_heading(doc, "5 ЗАКЛЮЧЕНИЕ")
    add_body(doc, "В результате выполнения лабораторных работ был разработан учебный, но практически ориентированный прототип серверного маркетплейса CloudFlow. Проект демонстрирует применение современных средств frontend- и backend-разработки, работу с типизированным API, организацию безопасной аутентификации, хранение приватных данных, построение ролевой модели и реализацию бизнес-процесса покупки серверной инфраструктуры. Структура решения позволяет расширять проект в дальнейшем, например, за счет реальной платежной интеграции, поддержки нескольких экземпляров товара, расширенной аналитики и более глубокого администрирования.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
